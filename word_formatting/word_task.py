from __future__ import annotations

import argparse
import csv
import json
import random
import shutil
import sys
import tempfile
import urllib.request
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_LINE_SPACING


SAMPLE_URL = (
    "https://huggingface.co/datasets/xlangai/ubuntu_osworld_file_cache/resolve/main/"
    "libreoffice_writer/0810415c-bde4-4443-9047-d5f70165a697/Novels_Intro_Packet.docx"
)

DOCX_NAME = "Novels_Intro_Packet.docx"
EDITED_DOCX_NAME = "Novels_Intro_Packet.edited.docx"

TASK_PROMPT = """Given Novels_Intro_Packet.docx, identify the first two body paragraphs, change only those two paragraphs to double line spacing, verify that all other text and formatting are unchanged, and generate a short Markdown formatting summary. No GUI, screenshot, external API, or cloud service should be used."""

SKILL_POOLS: dict[str, list[str]] = {
    "document_processing": [
        "word-docx",
        "office-document-specialist-suite",
        "file-converter",
        "all-to-markdown",
        "markdown-converter",
    ],
    "validation_audit": [
        "code-executor",
        "data-reconciliation-exceptions",
        "data-anomaly-detector",
        "data-analysis",
        "user-analysis-matrix",
    ],
    "summary_reporting": [
        "typora-visual-architect",
        "generate-report123",
        "markdown-converter",
        "sql-report-generator",
        "data2visualization",
    ],
}

SAMPLE_SEEDS = list(range(1, 11))


def unique_skills_from_pools(pools: dict[str, list[str]]) -> list[str]:
    seen: set[str] = set()
    skills: list[str] = []
    for pool_skills in pools.values():
        for skill in pool_skills:
            if skill not in seen:
                seen.add(skill)
                skills.append(skill)
    return skills


def sampled_skill_pools(seed: int, pools: dict[str, list[str]]) -> dict[str, list[str]]:
    rng = random.Random(seed)
    return {pool_name: [rng.choice(pool_skills)] for pool_name, pool_skills in pools.items()}


def build_conditions() -> dict[str, dict[str, Any]]:
    conditions: dict[str, dict[str, Any]] = {
        "no_skill": {
            "label": "No-skill baseline",
            "skills": [],
            "skill_pools": {},
            "condition_type": "baseline",
            "sample_seed": "",
            "note": "No skill document is provided. This measures direct model behavior.",
        },
        "multi_pool_all": {
            "label": "All candidates from every skill pool",
            "skills": unique_skills_from_pools(SKILL_POOLS),
            "skill_pools": SKILL_POOLS,
            "condition_type": "multi_pool_all",
            "sample_seed": "",
            "pool_samples": {},
            "note": "Provide every candidate skill from every pool and let the model choose and combine skill guidance.",
        },
    }

    for seed in SAMPLE_SEEDS:
        pools = sampled_skill_pools(seed, SKILL_POOLS)
        pool_samples = {pool_name: pool_skills[0] for pool_name, pool_skills in pools.items()}
        conditions[f"multi_pool_sample_s{seed:02d}"] = {
            "label": f"One sampled skill from each pool, seed {seed}",
            "skills": unique_skills_from_pools(pools),
            "skill_pools": pools,
            "condition_type": "multi_pool_sample",
            "sample_seed": seed,
            "note": (
                "Provide exactly one deterministic random skill from each pool. "
                "Compare repeated sampled combinations against the all-candidates selection condition."
            ),
            "pool_samples": pool_samples,
        }
    return conditions


CONDITIONS: dict[str, dict[str, Any]] = build_conditions()


@dataclass(frozen=True)
class Paths:
    root: Path
    data_dir: Path
    original_docx: Path
    target_meta: Path
    runs_dir: Path
    results_dir: Path
    results_csv: Path


def paths() -> Paths:
    root = Path(__file__).resolve().parent
    return Paths(
        root=root,
        data_dir=root / "data" / "original",
        original_docx=root / "data" / "original" / DOCX_NAME,
        target_meta=root / "data" / "target_paragraphs.json",
        runs_dir=root / "runs",
        results_dir=root / "results",
        results_csv=root / "results" / "results.csv",
    )


def skills_root() -> Path:
    return Path(__file__).resolve().parents[1] / "skills"


def ensure_dirs(p: Paths) -> None:
    p.data_dir.mkdir(parents=True, exist_ok=True)
    p.runs_dir.mkdir(parents=True, exist_ok=True)
    p.results_dir.mkdir(parents=True, exist_ok=True)
    for condition in CONDITIONS:
        (p.runs_dir / condition / "artifacts").mkdir(parents=True, exist_ok=True)


def download_sample(output_path: Path, force: bool = False) -> None:
    if output_path.exists() and not force:
        return

    output_path.parent.mkdir(parents=True, exist_ok=True)
    request = urllib.request.Request(
        SAMPLE_URL,
        headers={"User-Agent": "skill-composition-pilot/0.1"},
    )

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp_path = Path(tmp.name)

    try:
        with urllib.request.urlopen(request, timeout=120) as response:
            tmp_path.write_bytes(response.read())
        shutil.move(str(tmp_path), output_path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink()


def style_name(paragraph: Any) -> str:
    return paragraph.style.name if paragraph.style is not None else ""


def is_heading_like(paragraph: Any) -> bool:
    name = style_name(paragraph).lower()
    blocked = ("heading", "title", "subtitle", "toc")
    return any(part in name for part in blocked)


def candidate_body_paragraph_indices(doc: Any, min_chars: int = 20) -> list[int]:
    candidates: list[int] = []
    fallback: list[int] = []

    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        fallback.append(index)
        if is_heading_like(paragraph):
            continue
        if len(text) < min_chars:
            continue
        candidates.append(index)

    if len(candidates) >= 2:
        return candidates[:2]
    return fallback[:2]


def load_or_create_target_meta(p: Paths) -> list[dict[str, Any]]:
    if p.target_meta.exists():
        return json.loads(p.target_meta.read_text(encoding="utf-8"))

    doc = Document(str(p.original_docx))
    indices = candidate_body_paragraph_indices(doc)
    if len(indices) < 2:
        raise RuntimeError("Could not identify two body paragraphs in the original DOCX.")

    targets = [
        {
            "index": index,
            "style": style_name(doc.paragraphs[index]),
            "text": doc.paragraphs[index].text,
        }
        for index in indices
    ]
    p.target_meta.write_text(json.dumps(targets, ensure_ascii=False, indent=2), encoding="utf-8")
    return targets


def prompt_for_condition(condition: str, config: dict[str, Any]) -> str:
    skills = config["skills"]
    skills_text = "\n".join(f"- {skill}" for skill in skills) if skills else "- none"
    skill_pools = config.get("skill_pools", {})
    skill_pools_text = format_skill_pools(skill_pools)
    pool_samples_text = format_pool_samples(config.get("pool_samples", {}))
    skill_docs = load_skill_pool_docs(skill_pools) if skill_pools else load_skill_docs(skills)
    return f"""# Experiment Condition: {condition}

{config["label"]}

## Condition Metadata

- condition_type: {config.get("condition_type", "")}
- sample_seed: {config.get("sample_seed", "")}

## Skill Pools

{skill_pools_text}

## Pool Samples

{pool_samples_text}

## Task

{TASK_PROMPT}

## Available Skills

{skills_text}

## Available Skill Documents

{skill_docs}

## Condition Note

{config["note"]}

## Required Outputs

Save the edited file and summary as:

```text
artifacts/{EDITED_DOCX_NAME}
artifacts/summary.md
```

Do not use GUI operations, screenshots, external APIs, or cloud services.
"""


def load_skill_docs(skills: list[str]) -> str:
    if not skills:
        return "No external skill document is provided in this condition."

    blocks: list[str] = []
    for slug in skills:
        skill_md = skills_root() / slug / "SKILL.md"
        if skill_md.exists():
            content = skill_md.read_text(encoding="utf-8", errors="replace").strip()
            blocks.append(f"### {slug}\n\n```markdown\n{content}\n```")
        else:
            blocks.append(
                f"### {slug}\n\n"
                f"Skill document not downloaded yet. Expected path: {skill_md}"
            )
    return "\n\n".join(blocks)


def format_skill_pools(skill_pools: dict[str, list[str]]) -> str:
    if not skill_pools:
        return "- none"
    lines: list[str] = []
    for pool_name, skills in skill_pools.items():
        lines.append(f"### {pool_name}")
        for skill in skills:
            lines.append(f"- {skill}")
        lines.append("")
    return "\n".join(lines).strip()


def format_pool_samples(pool_samples: dict[str, str]) -> str:
    if not pool_samples:
        return "- none"
    return "\n".join(f"- {pool_name}: {skill}" for pool_name, skill in pool_samples.items())


def load_skill_pool_docs(skill_pools: dict[str, list[str]]) -> str:
    if not skill_pools:
        return "No external skill document is provided in this condition."

    blocks: list[str] = []
    for pool_name, skills in skill_pools.items():
        blocks.append(f"## Pool: {pool_name}")
        blocks.append(load_skill_docs(skills))
    return "\n\n".join(blocks)


def write_condition_prompts(p: Paths) -> None:
    manifest: dict[str, Any] = {}
    for condition, config in CONDITIONS.items():
        run_dir = p.runs_dir / condition
        (run_dir / "artifacts").mkdir(parents=True, exist_ok=True)
        (run_dir / "prompt.md").write_text(prompt_for_condition(condition, config), encoding="utf-8")
        (run_dir / "skills_given.txt").write_text(
            "\n".join(config["skills"]) if config["skills"] else "none\n",
            encoding="utf-8",
        )
        manifest[condition] = {
            "label": config["label"],
            "condition_type": config.get("condition_type", ""),
            "sample_seed": config.get("sample_seed", ""),
            "pool_samples": config.get("pool_samples", {}),
            "skill_pools": config.get("skill_pools", {}),
            "skills": config["skills"],
        }
    (p.runs_dir / "condition_manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def init_results_csv(p: Paths) -> None:
    if p.results_csv.exists():
        return
    with p.results_csv.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(
            f,
            fieldnames=[
                "timestamp",
                "condition",
                "success",
                "docx_exists",
                "summary_exists",
                "text_unchanged",
                "target_double_spaced",
                "other_formatting_unchanged",
                "issues",
            ],
        )
        writer.writeheader()


def prepare(args: argparse.Namespace) -> None:
    p = paths()
    ensure_dirs(p)
    download_sample(p.original_docx, force=args.force_download)
    targets = load_or_create_target_meta(p)
    write_condition_prompts(p)
    init_results_csv(p)
    print(f"Prepared Word pilot at {p.root}")
    print(f"Original DOCX: {p.original_docx}")
    print("Target paragraphs:")
    for target in targets:
        preview = target["text"].replace("\n", " ")[:120]
        print(f"  paragraph[{target['index']}] style={target['style']!r} text={preview!r}")


def normalize_length(value: Any) -> int | None:
    if value is None:
        return None
    return int(value)


def normalize_line_spacing(value: Any) -> str | float | int | None:
    if value is None:
        return None
    if isinstance(value, float):
        return round(value, 4)
    if isinstance(value, int):
        return value
    try:
        return int(value)
    except Exception:
        return str(value)


def normalize_enum(value: Any) -> str | None:
    return str(value) if value is not None else None


def color_value(run: Any) -> str | None:
    color = run.font.color
    if color is None or color.rgb is None:
        return None
    return str(color.rgb)


def font_size_value(run: Any) -> int | None:
    if run.font.size is None:
        return None
    return int(run.font.size)


def paragraph_signature(paragraph: Any, include_line_spacing: bool = True) -> dict[str, Any]:
    fmt = paragraph.paragraph_format
    signature: dict[str, Any] = {
        "text": paragraph.text,
        "style": style_name(paragraph),
        "alignment": normalize_enum(fmt.alignment),
        "left_indent": normalize_length(fmt.left_indent),
        "right_indent": normalize_length(fmt.right_indent),
        "first_line_indent": normalize_length(fmt.first_line_indent),
        "space_before": normalize_length(fmt.space_before),
        "space_after": normalize_length(fmt.space_after),
        "runs": [
            {
                "text": run.text,
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
                "font_name": run.font.name,
                "font_size": font_size_value(run),
                "font_color": color_value(run),
            }
            for run in paragraph.runs
        ],
    }
    if include_line_spacing:
        signature["line_spacing"] = normalize_line_spacing(fmt.line_spacing)
        signature["line_spacing_rule"] = normalize_enum(fmt.line_spacing_rule)
    return signature


def is_double_spaced(paragraph: Any) -> bool:
    fmt = paragraph.paragraph_format
    if fmt.line_spacing_rule == WD_LINE_SPACING.DOUBLE:
        return True
    value = fmt.line_spacing
    if isinstance(value, float) and abs(value - 2.0) < 0.01:
        return True
    return False


def run_oracle(args: argparse.Namespace) -> None:
    p = paths()
    ensure_dirs(p)
    if not p.original_docx.exists():
        download_sample(p.original_docx)
    targets = load_or_create_target_meta(p)
    target_indices = {target["index"] for target in targets}

    doc = Document(str(p.original_docx))
    for index in target_indices:
        paragraph = doc.paragraphs[index]
        paragraph.paragraph_format.line_spacing = 2.0
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    output_dir = p.runs_dir / args.condition / "artifacts"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_docx = output_dir / EDITED_DOCX_NAME
    summary_md = output_dir / "summary.md"
    doc.save(str(output_docx))
    summary_md.write_text(
        "# Formatting Summary\n\n"
        f"- Updated only the first two detected body paragraphs to double line spacing.\n"
        f"- Target paragraph indices: {', '.join(str(i) for i in sorted(target_indices))}.\n"
        "- Text content and non-target paragraph formatting were left unchanged.\n",
        encoding="utf-8",
    )
    print(f"Oracle output written to {output_docx}")
    print(f"Oracle summary written to {summary_md}")


def verify_paths_for_condition(condition: str, output_docx: str | None, summary: str | None) -> tuple[Path, Path]:
    p = paths()
    default_artifacts = p.runs_dir / condition / "artifacts"
    return (
        Path(output_docx) if output_docx else default_artifacts / EDITED_DOCX_NAME,
        Path(summary) if summary else default_artifacts / "summary.md",
    )


def verify_output(condition: str, output_docx: Path, summary_md: Path) -> dict[str, Any]:
    p = paths()
    issues: list[str] = []
    checks = {
        "docx_exists": output_docx.exists(),
        "summary_exists": summary_md.exists() and summary_md.stat().st_size > 0,
        "text_unchanged": False,
        "target_double_spaced": False,
        "other_formatting_unchanged": False,
    }

    if not checks["docx_exists"]:
        issues.append(f"Missing edited DOCX: {output_docx}")
    if not checks["summary_exists"]:
        issues.append(f"Missing or empty summary.md: {summary_md}")
    if not output_docx.exists():
        return finish_result(condition, checks, issues)

    try:
        original = Document(str(p.original_docx))
        edited = Document(str(output_docx))
    except Exception as exc:
        issues.append(f"Could not open DOCX: {exc}")
        return finish_result(condition, checks, issues)

    targets = load_or_create_target_meta(p)
    target_indices = {target["index"] for target in targets}

    original_text = [paragraph.text for paragraph in original.paragraphs]
    edited_text = [paragraph.text for paragraph in edited.paragraphs]
    checks["text_unchanged"] = original_text == edited_text
    if not checks["text_unchanged"]:
        issues.append("Paragraph text or paragraph count changed.")

    target_results = []
    for index in target_indices:
        if index >= len(edited.paragraphs):
            target_results.append(False)
            issues.append(f"Target paragraph index {index} is missing in edited DOCX.")
            continue
        target_results.append(is_double_spaced(edited.paragraphs[index]))
        if not target_results[-1]:
            issues.append(f"Target paragraph index {index} is not double spaced.")
    checks["target_double_spaced"] = bool(target_results) and all(target_results)

    non_target_ok = True
    count = min(len(original.paragraphs), len(edited.paragraphs))
    for index in range(count):
        include_line_spacing = index not in target_indices
        original_sig = paragraph_signature(original.paragraphs[index], include_line_spacing=include_line_spacing)
        edited_sig = paragraph_signature(edited.paragraphs[index], include_line_spacing=include_line_spacing)
        if original_sig != edited_sig:
            non_target_ok = False
            if index in target_indices:
                issues.append(f"Target paragraph index {index} changed beyond line spacing.")
            else:
                issues.append(f"Non-target paragraph index {index} formatting/text changed.")
            if len(issues) >= 12:
                issues.append("Additional formatting differences omitted.")
                break
    if len(original.paragraphs) != len(edited.paragraphs):
        non_target_ok = False
    checks["other_formatting_unchanged"] = non_target_ok

    return finish_result(condition, checks, issues)


def finish_result(condition: str, checks: dict[str, bool], issues: list[str]) -> dict[str, Any]:
    success = all(checks.values())
    return {
        "timestamp": datetime.now().isoformat(timespec="seconds"),
        "condition": condition,
        "success": success,
        **checks,
        "issues": " | ".join(issues),
    }


def append_result(result: dict[str, Any]) -> None:
    p = paths()
    init_results_csv(p)
    with p.results_csv.open("a", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(
            f,
            fieldnames=[
                "timestamp",
                "condition",
                "success",
                "docx_exists",
                "summary_exists",
                "text_unchanged",
                "target_double_spaced",
                "other_formatting_unchanged",
                "issues",
            ],
        )
        writer.writerow(result)


def print_result(result: dict[str, Any]) -> None:
    print(json.dumps(result, ensure_ascii=False, indent=2))


def verify(args: argparse.Namespace) -> None:
    output_docx, summary_md = verify_paths_for_condition(args.condition, args.output_docx, args.summary)
    result = verify_output(args.condition, output_docx, summary_md)
    append_result(result)
    print_result(result)
    if not result["success"]:
        raise SystemExit(1)


def verify_all(args: argparse.Namespace) -> None:
    failed = False
    for condition in CONDITIONS:
        output_docx, summary_md = verify_paths_for_condition(condition, None, None)
        result = verify_output(condition, output_docx, summary_md)
        append_result(result)
        print_result(result)
        failed = failed or not result["success"]
    if failed:
        raise SystemExit(1)


def show_targets(args: argparse.Namespace) -> None:
    p = paths()
    if not p.original_docx.exists():
        download_sample(p.original_docx)
    targets = load_or_create_target_meta(p)
    print(json.dumps(targets, ensure_ascii=False, indent=2))


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Word formatting pilot harness.")
    subparsers = parser.add_subparsers(dest="command", required=True)

    prepare_parser = subparsers.add_parser("prepare", help="Download sample, create prompts, init results.")
    prepare_parser.add_argument("--force-download", action="store_true", help="Re-download the sample DOCX.")
    prepare_parser.set_defaults(func=prepare)

    oracle_parser = subparsers.add_parser("oracle", help="Generate a known-correct output for one condition.")
    oracle_parser.add_argument("--condition", choices=sorted(CONDITIONS), default="multi_pool_sample_s02")
    oracle_parser.set_defaults(func=run_oracle)

    verify_parser = subparsers.add_parser("verify", help="Verify one condition output.")
    verify_parser.add_argument("--condition", choices=sorted(CONDITIONS), required=True)
    verify_parser.add_argument("--output-docx", help="Optional edited DOCX path. Defaults to runs/<condition>/artifacts.")
    verify_parser.add_argument("--summary", help="Optional summary.md path. Defaults to runs/<condition>/artifacts.")
    verify_parser.set_defaults(func=verify)

    verify_all_parser = subparsers.add_parser("verify-all", help="Verify all condition folders.")
    verify_all_parser.set_defaults(func=verify_all)

    show_targets_parser = subparsers.add_parser("show-targets", help="Print selected body paragraph metadata.")
    show_targets_parser.set_defaults(func=show_targets)

    return parser


def main(argv: list[str] | None = None) -> int:
    parser = build_parser()
    args = parser.parse_args(argv)
    args.func(args)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
